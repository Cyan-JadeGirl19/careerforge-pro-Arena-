"""Document export: plain text, JSON, DOCX, and PDF.

All layouts are single-column and parser-safe (no tables, text boxes, or
graphics), per the ATS Enterprise agreement. The Modern layout adds only
subtle typographic hierarchy.
"""
import io
import json

from .content import CvContent

_BRAND = (0x5B, 0x4B, 0xDB)


def to_plain_text(c: CvContent) -> str:
    lines: list[str] = []
    lines.append(c.name.upper())
    contact = "  ·  ".join(x for x in [c.email, c.phone, c.location, *c.links] if x)
    if contact:
        lines.append(contact)
    if c.headline:
        lines.append("")
        lines.append(c.headline)
    if c.summary:
        lines += ["", "SUMMARY", c.summary]

    def section(title: str, items: list[str]) -> None:
        nonlocal lines
        if items:
            lines.append("")
            lines.append(title.upper())
            lines.extend(f"- {i}" for i in items)

    if c.skills:
        lines += ["", "SKILLS", ", ".join(c.skills)]
    if c.experience:
        lines += ["", "EXPERIENCE"]
        for e in c.experience:
            head = " · ".join(x for x in [e.title, e.company, e.dates] if x)
            lines.append(head)
            lines += [f"- {b}" for b in e.bullets]
            lines.append("")
    section("EDUCATION", [
        " · ".join(x for x in [e.degree, e.institution, e.year] if x)
        for e in c.education
    ])
    section("CERTIFICATIONS", c.certifications)
    section("PROJECTS", c.projects)
    section("LANGUAGES", c.languages)
    return "\n".join(lines).strip() + "\n"


def to_json(c: CvContent) -> str:
    return json.dumps(c.to_dict(), indent=2, ensure_ascii=False) + "\n"


def to_docx(c: CvContent) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    r = p.add_run(c.name)
    r.bold = True
    r.font.size = Pt(16)

    contact = " · ".join(x for x in [c.email, c.phone, c.location, *c.links] if x)
    if contact:
        cp = doc.add_paragraph()
        cp.add_run(contact).font.size = Pt(9)

    if c.headline:
        hp = doc.add_paragraph()
        hr = hp.add_run(c.headline)
        hr.bold = True
        hr.font.size = Pt(11)
        if c.layout != "ats_single_column":
            hr.font.color.rgb = RGBColor(*_BRAND)

    def heading(text: str) -> None:
        hp = doc.add_paragraph()
        hr = hp.add_run(text.upper())
        hr.bold = True
        hr.font.size = Pt(11)
        if c.layout == "ats_single_column":
            hr.font.color.rgb = RGBColor(0, 0, 0)

    if c.summary:
        heading("Summary")
        doc.add_paragraph(c.summary)

    if c.skills:
        heading("Skills")
        doc.add_paragraph(", ".join(c.skills))

    if c.experience:
        heading("Experience")
        for e in c.experience:
            ep = doc.add_paragraph()
            er = ep.add_run(" · ".join(x for x in [e.title, e.company, e.dates] if x))
            er.bold = True
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if c.education:
        heading("Education")
        for e in c.education:
            doc.add_paragraph(
                " · ".join(x for x in [e.degree, e.institution, e.year] if x),
                style="List Bullet",
            )

    for title, items in (
        ("Certifications", c.certifications),
        ("Projects", c.projects),
        ("Languages", c.languages),
    ):
        if items:
            heading(title)
            for i in items:
                doc.add_paragraph(i, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(c: CvContent) -> bytes:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    name_style = ParagraphStyle("name", fontSize=18, leading=21, spaceAfter=2)
    contact_style = ParagraphStyle("contact", fontSize=8.5, leading=11, textColor=HexColor("#555a68"))
    headline_style = ParagraphStyle("headline", fontSize=11, leading=14, spaceAfter=8, textColor=HexColor("#%02x%02x%02x" % _BRAND))
    section_style = ParagraphStyle("section", fontSize=11, leading=14, spaceBefore=10, spaceAfter=3, textColor=HexColor("#000000"))
    body_style = ParagraphStyle("body", fontSize=9.5, leading=13)
    bullet_style = ParagraphStyle("bullet", fontSize=9.5, leading=13, leftIndent=10)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        title=c.name or "CV",
        author="CareerForge Pro",
    )
    story: list = []
    if c.name:
        story.append(Paragraph(c.name, name_style))
    contact = " · ".join(x for x in [c.email, c.phone, c.location, *c.links] if x)
    if contact:
        story.append(Paragraph(contact, contact_style))
    if c.headline:
        story.append(Paragraph(c.headline, headline_style))

    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    if c.summary:
        story.append(Paragraph("Summary", section_style))
        story.append(Paragraph(esc(c.summary), body_style))

    if c.skills:
        story.append(Paragraph("Skills", section_style))
        story.append(Paragraph(esc(", ".join(c.skills)), body_style))

    if c.experience:
        story.append(Paragraph("Experience", section_style))
        for e in c.experience:
            story.append(
                Paragraph(
                    f"<b>{esc(' · '.join(x for x in [e.title, e.company, e.dates] if x))}</b>",
                    body_style,
                )
            )
            if e.bullets:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(esc(b), bullet_style), leftIndent=6) for b in e.bullets],
                        bulletType="bullet",
                        start="•",
                        leftIndent=12,
                    )
                )
            story.append(Spacer(1, 4))

    def list_section(title: str, items: list[str]) -> None:
        if not items:
            return
        story.append(Paragraph(title, section_style))
        story.append(
            ListFlowable(
                [ListItem(Paragraph(esc(i), bullet_style), leftIndent=6) for i in items],
                bulletType="bullet",
                start="•",
                leftIndent=12,
            )
        )

    list_section("Education", [
        " · ".join(x for x in [e.degree, e.institution, e.year] if x)
        for e in c.education
    ])
    list_section("Certifications", c.certifications)
    list_section("Projects", c.projects)
    list_section("Languages", c.languages)

    doc.build(story)
    return buf.getvalue()


EXPORTERS = {
    "txt": lambda c: to_plain_text(c).encode("utf-8"),
    "json": lambda c: to_json(c).encode("utf-8"),
    "docx": to_docx,
    "pdf": to_pdf,
}

MEDIA_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
