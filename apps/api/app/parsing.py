"""CV intake: file text extraction + structured parsing.

Everything extracted is traceable to the candidate's own source text —
the parser never invents content. Fields it cannot determine
confidently are reported in ``extraction_notes`` so the candidate can
review and confirm (agreed CV-first intake flow, step 2).
"""
import io
import re
from dataclasses import dataclass, field

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?<![\w.])(\+?\d[\d\s()-]{8,}\d)(?![\w])")
LINK_RE = re.compile(r"(?:https?://|www\.)[^\s(),;]+|(?:linkedin\.com|github\.com)/[\w/=-]+")
BULLET_RE = re.compile(r"^\s*[-\u2022*]\s*")
DATE_RANGE_RE = re.compile(
    r"\b((?:19|20)\d{2})\s*(?:-|–|—|to)\s*((?:19|20)\d{2}|present|current|now|today)\b",
    re.IGNORECASE,
)
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")

SECTION_KEYWORDS = {
    "summary": (
        "professional summary", "profile summary", "summary", "profile",
        "objective", "career objective", "personal statement", "about me",
        "introduction",
    ),
    "experience": (
        "work experience", "professional experience", "employment history",
        "employment", "work history", "career history", "experience",
        "relevant experience",
    ),
    "education": ("education", "academic background", "qualifications"),
    "skills": (
        "core skills", "key skills", "technical skills", "skill set",
        "skills", "competencies",
    ),
    "certifications": (
        "certifications", "certificates", "licenses",
        "professional development", "licences",
    ),
    "projects": ("key projects", "personal projects", "projects"),
    "languages": ("languages spoken", "languages"),
}

_TITLE_WORDS = (
    "manager", "engineer", "analyst", "agent", "lead", "specialist",
    "officer", "director", "consultant", "developer", "designer",
    "executive", "coordinator", "administrator", "supervisor",
    "coordinator", "representative", "associate", "intern", "head",
    "chief", "founder", "owner", "practitioner", "advisor", "writer",
    "support", "success",
)


@dataclass
class ParsedCv:
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    links: list[str] = field(default_factory=list)
    summary: str = ""
    experience: list[dict] = field(default_factory=list)
    education: list[dict] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    projects: list[str] = field(default_factory=list)
    languages: list[str] = field(default_factory=list)
    other_sections: dict[str, list[str]] = field(default_factory=dict)
    extraction_notes: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "location": self.location,
            "links": self.links,
            "summary": self.summary,
            "experience": self.experience,
            "education": self.education,
            "skills": self.skills,
            "certifications": self.certifications,
            "projects": self.projects,
            "languages": self.languages,
            "other_sections": self.other_sections,
            "extraction_notes": self.extraction_notes,
        }

    def content_kwargs(self) -> dict:
        """Keys usable when constructing a CvContent (factual only)."""
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "location": self.location,
            "links": self.links,
        }


def extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def extract_file_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf_text(data)
    if lower.endswith(".docx"):
        return extract_docx_text(data)
    if lower.endswith((".txt", ".md", ".text")):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")


def _is_heading(line: str) -> str | None:
    stripped = line.strip().strip(":-|").strip().lower()
    if not stripped or len(stripped) > 40:
        return None
    for section, keywords in SECTION_KEYWORDS.items():
        if stripped in keywords or any(stripped == k for k in keywords):
            return section
    return None


def _likely_name(line: str) -> bool:
    if not line or len(line) > 50:
        return False
    if "@" in line or re.search(r"\d", line):
        return False
    words = line.split()
    if 1 <= len(words) <= 4:
        return any(w and w[0].isupper() for w in words)
    return False


def _split_role_line(line: str) -> dict:
    """Split 'Title, Company (dates)' style lines."""
    dates = ""
    m = DATE_RANGE_RE.search(line)
    if m:
        dates = f"{m.group(1)} - {m.group(2)}"
        line = line[: m.start()] + line[m.end():]
    line = re.sub(r"\(\s*\)", "", line)  # leftover empty parentheses
    line = line.strip(" \t|-–—()")
    title, company = line, ""
    for sep in (" | ", " – ", " — "):
        if sep in line:
            left, right = line.split(sep, 1)
            if _is_title_like(left):
                title, company = left.strip(), right.strip()
            else:
                title, company = right.strip(), left.strip()
            break
    else:
        if "," in line:
            left, right = line.split(",", 1)
            if len(right.strip()) <= 60 and not left.strip().startswith("http"):
                title, company = left.strip(), right.strip()
    return {"title": title, "company": company, "dates": dates, "bullets": []}


def _is_title_like(text: str) -> bool:
    return any(w in text.lower() for w in _TITLE_WORDS)


def parse_cv_text(text: str) -> ParsedCv:
    parsed = ParsedCv()
    lines = [ln.rstrip() for ln in text.splitlines()]
    nonempty = [ln.strip() for ln in lines if ln.strip()]

    # --- contact ---------------------------------------------------------
    blob = "\n".join(nonempty[:12])
    em = EMAIL_RE.search(blob)
    if em:
        parsed.email = em.group(0)
    else:
        parsed.extraction_notes.append("No email address found — please add one.")
    ph = PHONE_RE.search("\n".join(nonempty))
    if ph:
        parsed.phone = ph.group(1).strip()
    else:
        parsed.extraction_notes.append("No phone number found — optional, but recommended.")
    for ln in nonempty[:10]:
        lm = LINK_RE.search(ln)
        if lm:
            parsed.links.append(lm.group(0))
    # location: short line near the top containing a comma + place words
    for ln in nonempty[:8]:
        if ln.lower().strip() in ("location",):
            continue
        if re.search(r"\b(johannesburg|cape town|pretoria|durban|sandton|bloemfontein|south africa)\b", ln, re.IGNORECASE):
            parsed.location = ln.strip()
            break
    if not parsed.location:
        parsed.extraction_notes.append("Location not detected — please confirm it.")

    # --- name -------------------------------------------------------------
    for ln in nonempty[:8]:
        if _likely_name(ln) and not _is_heading(ln) and not EMAIL_RE.search(ln):
            parsed.name = ln.strip()
            break
    if not parsed.name:
        parsed.extraction_notes.append(
            "Could not determine your name with confidence — please review."
        )

    # --- sections ----------------------------------------------------------
    current: str | None = None
    summary_lines: list[str] = []
    for ln in nonempty:
        heading = _is_heading(ln)
        if heading:
            current = heading
            continue
        bullet = BULLET_RE.match(ln)
        if current == "experience":
            if bullet:
                if parsed.experience:
                    parsed.experience[-1]["bullets"].append(BULLET_RE.sub("", ln).strip())
            else:
                entry = _split_role_line(ln)
                if entry["title"] or entry["company"]:
                    parsed.experience.append(entry)
        elif current == "summary":
            summary_lines.append(ln)
        elif current == "education":
            entry = _split_role_line(ln)
            if entry["title"] or entry["company"]:
                entry["bullets"] = []
                parsed.education.append(
                    {
                        "degree": entry["title"],
                        "institution": entry["company"],
                        "year": entry["dates"],
                    }
                )
        elif current in ("skills", "certifications", "projects", "languages"):
            for piece in re.split(r"[,;\n]", ln):
                piece = piece.strip().strip("-\u2022* ").strip()
                if piece and len(piece) > 1:
                    getattr(parsed, current).append(piece)
        elif current is None:
            if re.search(r"\b(johannesburg|cape town|pretoria|durban|south africa)\b", ln, re.IGNORECASE) and not parsed.location:
                parsed.location = ln.strip()

    parsed.summary = " ".join(summary_lines).strip()
    if not parsed.summary:
        parsed.extraction_notes.append(
            "No summary section found — one will be drafted from your experience for your approval."
        )
    if not parsed.experience:
        parsed.extraction_notes.append(
            "Work history was not recognised in a standard format — please review and confirm the entries."
        )
    if not parsed.skills:
        parsed.extraction_notes.append(
            "No skills section found — skills will be inferred from your experience for your confirmation."
        )
    return parsed
