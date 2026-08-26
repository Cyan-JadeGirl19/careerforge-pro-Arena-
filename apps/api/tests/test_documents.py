"""Document engine tests: upload/parse, versions, tailoring, export."""
import io
import json
import re

API = "/api/v1"

FULL_CV = """\
Thando Ndlovu
thando.ndlovu@example.com
+27 82 123 4567
Johannesburg, South Africa
linkedin.com/in/thandon

Summary
Customer operations professional with 6 years of experience supporting remote SaaS teams.

Experience
Support Team Lead, Luno (2022 - present)
- Led a remote team of 4 support agents; raised CSAT from 82% to 91%
- Handled 40+ tickets per day across email and chat
- Reduced repeat tickets by 22% with a new knowledge base
Support Agent, PayFast (2019 - 2022)
- Resolved 25+ customer issues daily while meeting SLA targets
- Built onboarding guides that cut new-agent ramp-up time from 6 weeks to 3

Education
BCom, University of Pretoria (2015 - 2018)

Skills
Project management, data analysis, communication, stakeholder management, Excel, customer success, remote collaboration
"""

JD = """\
Customer Success Manager — Remote (Africa)

We are looking for a Customer Success Manager to own onboarding and
retention for a growing portfolio of customers.

Requirements:
- 5+ years in customer success or customer service roles
- Strong data analysis skills; comfortable with Excel and dashboards
- Experience with stakeholder management across time zones
- Confident with python or similar for light automation (nice to have)
- Excellent communication and written communication skills
- Experience with CRM tools and renewal processes
"""


def _make_pdf(text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    style = getSampleStyleSheet()["Normal"]
    doc.build([Paragraph(line.replace("&", "&amp;"), style) for line in text.splitlines()])
    return buf.getvalue()


def _make_docx(text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- upload + parse -----------------------------------------------------------


def test_upload_pdf_is_parsed(client, consented_profile):
    res = client.post(
        f"{API}/profiles/{consented_profile}/cvs/upload",
        files={"file": ("cv.pdf", _make_pdf(FULL_CV), "application/pdf")},
    )
    assert res.status_code == 201
    body = res.json()
    assert body["source_type"] == "upload"
    parsed = body["parsed"]
    assert parsed["name"] == "Thando Ndlovu"
    assert parsed["email"] == "thando.ndlovu@example.com"
    assert parsed["location"].startswith("Johannesburg")
    assert len(parsed["experience"]) == 2
    assert parsed["experience"][0]["company"] == "Luno"
    assert any("CSAT from 82% to 91%" in b for b in parsed["experience"][0]["bullets"])
    assert "data analysis" in parsed["skills"]
    assert parsed["education"][0]["institution"] == "University of Pretoria"


def test_upload_docx_is_parsed(client, consented_profile):
    res = client.post(
        f"{API}/profiles/{consented_profile}/cvs/upload",
        files={"file": ("cv.docx", _make_docx(FULL_CV),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert res.status_code == 201
    assert res.json()["parsed"]["name"] == "Thando Ndlovu"


def test_upload_unsupported_type_rejected(client, consented_profile):
    res = client.post(
        f"{API}/profiles/{consented_profile}/cvs/upload",
        files={"file": ("cv.exe", b"MZ...", "application/octet-stream")},
    )
    assert res.status_code == 415
    assert res.json()["detail"]["code"] == "UNSUPPORTED_FILE"


def test_upload_requires_consent(client, profile_id):
    res = client.post(
        f"{API}/profiles/{profile_id}/cvs/upload",
        files={"file": ("cv.txt", FULL_CV.encode(), "text/plain")},
    )
    assert res.status_code == 409


def test_parsed_endpoint_for_pasted_cv(client, cv_id):
    res = client.get(f"{API}/cvs/{cv_id}/parsed")
    assert res.status_code == 200
    assert res.json()["name"] == "Thando Ndlovu"


# --- master + custom versions --------------------------------------------------


def test_build_masters_creates_one_per_top_role(client, cv_id):
    res = client.post(f"{API}/cvs/{cv_id}/versions/build-masters", json={})
    assert res.status_code == 201
    versions = res.json()
    assert len(versions) == 3
    # every master is role-focused (no style-only masters)
    assert {v["kind"] for v in versions} == {"master_role"}
    roles = [v["content"]["role_focus"] for v in versions]
    assert len({r.lower() for r in roles}) == 3, "three distinct roles"
    assert all(r for r in roles), "every master has a role focus"
    # all parser-safe single-column layouts
    assert {v["content"]["layout"] for v in versions} == {"role_specialist"}
    # titles carry the role
    for v in versions:
        assert v["title"].startswith("Master CV —")
    # factual integrity: name/email carried from source, nothing invented
    for v in versions:
        assert v["content"]["name"] == "Thando Ndlovu"
        assert v["content"]["source_profile_version"] == cv_id
        assert v["content"]["generation_timestamp"]


def test_build_masters_pins_requested_role_first(client, cv_id):
    res = client.post(
        f"{API}/cvs/{cv_id}/versions/build-masters",
        json={"role_focus": "marketing"},
    )
    assert res.status_code == 201
    versions = res.json()
    assert versions[0]["content"]["role_focus"] == "marketing"
    assert len(versions) == 3


def test_role_master_reorders_toward_role(client, cv_id):
    res = client.post(
        f"{API}/cvs/{cv_id}/versions",
        json={"kind": "master_role", "role_focus": "data analysis"},
    )
    assert res.status_code == 201
    skills = res.json()["content"]["skills"]
    # "data analysis" is in the candidate's real skills -> moved to the front
    assert skills[0].lower() == "data analysis"


def test_role_master_requires_focus(client, cv_id):
    res = client.post(f"{API}/cvs/{cv_id}/versions", json={"kind": "master_role"})
    assert res.status_code == 422
    assert res.json()["detail"]["code"] == "ROLE_FOCUS_REQUIRED"


def test_custom_version_emphasize_and_exclude(client, cv_id):
    res = client.post(
        f"{API}/cvs/{cv_id}/versions",
        json={
            "kind": "custom",
            "role_focus": "marketing",
            "emphasize": ["communication", "data analysis"],
            "exclude": ["onboarding"],
        },
    )
    assert res.status_code == 201
    content = res.json()["content"]
    assert content["role_focus"] == "marketing"
    skills = [s.lower() for s in content["skills"]]
    assert skills[0] == "communication"
    # excluded bullet removed everywhere
    for e in content["experience"]:
        assert not any("onboarding" in b.lower() for b in e["bullets"])


def test_custom_marketing_repositions_honestly(client, cv_id):
    """The sample CV has no marketing skills: the marketing version must
    still visibly target the role (headline + pivot summary) and explain
    transparently that no marketing-specific evidence was found."""
    res = client.post(
        f"{API}/cvs/{cv_id}/versions",
        json={"kind": "custom", "role_focus": "marketing"},
    )
    assert res.status_code == 201, res.text
    v = res.json()
    content = v["content"]
    # role-targeted headline, factual pivot summary (no invented experience)
    assert content["headline"].lower().startswith("marketing")
    assert "marketing" in content["summary"].lower()
    assert "moving into marketing" in content["summary"].lower()
    assert "Support Team Lead" in content["summary"]  # their real latest title
    # transparent notes: nothing matched + emphasised recorded
    assert v["notes"], "custom version must carry transparent notes"
    assert any("No marketing-specific skills" in n for n in v["notes"])
    # name/email still carried from the source, nothing invented
    assert content["name"] == "Thando Ndlovu"


def test_custom_role_with_matching_skills(client, cv_id):
    """'operations' keywords genuinely exist in the sample CV: the version
    should bring them forward and say so in the notes."""
    res = client.post(
        f"{API}/cvs/{cv_id}/versions",
        json={"kind": "custom", "role_focus": "operations"},
    )
    assert res.status_code == 201, res.text
    v = res.json()
    content = v["content"]
    assert content["headline"].lower().startswith("operations")
    assert "bringing" in content["summary"].lower()
    assert any(n.startswith("Skills matched to operations") for n in v["notes"])
    skills = [s.lower() for s in content["skills"]]
    assert skills[0] in ("operations", "data analysis", "stakeholder management")


# --- tailoring -----------------------------------------------------------------


def test_tailor_reports_coverage_and_gaps(client, consented_profile, cv_id):
    version = client.post(
        f"{API}/cvs/{cv_id}/versions", json={"kind": "master_modern"}
    ).json()
    jd = client.post(
        f"{API}/profiles/{consented_profile}/job-descriptions",
        json={"title": "Customer Success Manager", "company": "Acme", "text": JD},
    )
    assert jd.status_code == 201
    res = client.post(
        f"{API}/cv-versions/{version['id']}/tailor", json={"jd_id": jd.json()["id"]}
    )
    assert res.status_code == 201
    body = res.json()
    report = body["report"]
    kw_map = {k["keyword"]: k["in_candidate_profile"] for k in report["keywords"]}
    # genuinely supported -> present
    assert kw_map.get("data analysis") is True
    assert kw_map.get("customer success") is True
    # not in the candidate's profile -> gap, never invented
    assert kw_map.get("customer service") is False
    assert any("customer service" in g for g in report["gaps"])
    # every flagged gap corresponds to a keyword the profile does not support
    for g in report["gaps"]:
        kw = g.split("'")[1]
        assert kw_map.get(kw) is False
    assert body["content"]["job_description_version"] == jd.json()["id"]
    assert 0 < report["coverage"] < 100


def test_jd_requires_job_matching_consent(client, profile_id):
    res = client.post(
        f"{API}/profiles/{profile_id}/job-descriptions",
        json={"title": "Some Role", "text": JD},
    )
    assert res.status_code == 409


# --- export ---------------------------------------------------------------------


def _get_first_version(client, cv_id):
    return client.post(
        f"{API}/cvs/{cv_id}/versions/build-masters", json={}
    ).json()[0]


def test_export_docx_roundtrip(client, cv_id):
    from docx import Document

    v = _get_first_version(client, cv_id)
    res = client.get(f"{API}/cv-versions/{v['id']}/export", params={"format": "docx"})
    assert res.status_code == 200
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    doc = Document(io.BytesIO(res.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "THANDO NDLOVU" in text or "Thando Ndlovu" in text
    assert "data analysis" in text.lower()


def test_export_pdf_is_valid(client, cv_id):
    from pypdf import PdfReader

    v = _get_first_version(client, cv_id)
    res = client.get(f"{API}/cv-versions/{v['id']}/export", params={"format": "pdf"})
    assert res.status_code == 200
    assert res.content[:4] == b"%PDF"
    reader = PdfReader(io.BytesIO(res.content))
    assert len(reader.pages) >= 1
    page_text = reader.pages[0].extract_text() or ""
    assert "NDLOVU" in page_text.upper()


def test_export_txt_and_json(client, cv_id):
    v = _get_first_version(client, cv_id)
    txt = client.get(f"{API}/cv-versions/{v['id']}/export", params={"format": "txt"})
    assert txt.status_code == 200
    assert "THANDO NDLOVU" in txt.text
    j = client.get(f"{API}/cv-versions/{v['id']}/export", params={"format": "json"})
    assert j.status_code == 200
    body = json.loads(j.text)
    assert body["name"] == "Thando Ndlovu"
    assert body["skills"]


def test_tailored_export(client, consented_profile, cv_id):
    version = client.post(
        f"{API}/cvs/{cv_id}/versions", json={"kind": "master_ats"}
    ).json()
    jd = client.post(
        f"{API}/profiles/{consented_profile}/job-descriptions",
        json={"title": "Customer Success Manager", "text": JD},
    ).json()
    tailored = client.post(
        f"{API}/cv-versions/{version['id']}/tailor", json={"jd_id": jd["id"]}
    ).json()
    res = client.get(f"{API}/tailored/{tailored['id']}/export", params={"format": "pdf"})
    assert res.status_code == 200
    assert res.content[:4] == b"%PDF"


# --- tailoring rewrite engine -----------------------------------------------------

from app.builders import (  # noqa: E402
    KEYWORD_BRIDGES,
    _relabel_bullet,
    _role_relevance,
    _targeted_summary,
)


def test_relabel_bullet_preserves_facts():
    b, kw = _relabel_bullet(
        "Resolved 25+ customer issues daily while meeting SLA targets",
        ["customer support"],
    )
    assert kw == "customer support"
    assert b.startswith("Customer Support: ")
    assert "25+" in b and "SLA targets" in b  # original fact intact


def test_relabel_bullet_skips_already_present_keyword():
    b, kw = _relabel_bullet("Provided customer support for 40 customers", ["customer support"])
    assert kw is None and b == "Provided customer support for 40 customers"


def test_relabel_bullet_no_bridge_no_change():
    b, kw = _relabel_bullet("Managed the office coffee machine", ["quantum physics"])
    assert kw is None
    assert b == "Managed the office coffee machine"


def test_role_relevance_ranks_relevant_role_higher():
    kws = ["customer support", "onboarding"]
    a = {"title": "Support Lead", "bullets": ["handled customer support tickets and onboarding"]}
    b = {"title": "Warehouse Clerk", "bullets": ["stacked boxes"]}
    assert _role_relevance(a, kws) > _role_relevance(b, kws)


def test_targeted_summary_uses_only_supported_keywords():
    # "customer support" is in the corpus and not part of the job title,
    # so it qualifies; "retention" is unsupported and must not be claimed.
    s, added = _targeted_summary(
        "Support lead.",
        "Customer Success Manager",
        ["retention", "customer support"],
        "i drive customer support and csat",
    )
    assert [k.lower() for k in added] == ["customer support"]
    assert "retention" not in [k.lower() for k in added], "unsupported keyword must not be claimed"
    assert "Customer Success Manager" in s


def test_targeted_summary_skips_keywords_inside_job_title():
    # "customer success" is inside "Customer Success Manager" - restating
    # it next to the title is noise, so nothing is added.
    s, added = _targeted_summary(
        "Support lead.",
        "Customer Success Manager",
        ["customer success"],
        "i drive customer success",
    )
    assert added == []
    assert s == "Support lead."


def test_tailor_rewrites_summary_and_bullets(client, consented_profile, cv_id):
    version = client.post(
        f"{API}/cvs/{cv_id}/versions",
        json={"kind": "master_role", "role_focus": "customer success"},
    ).json()
    jd = client.post(
        f"{API}/profiles/{consented_profile}/job-descriptions",
        json={
            "title": "Customer Success Manager",
            "company": "Acme",
            "text": (
                "We need a customer success manager to drive customer success, "
                "retention and churn reduction for SaaS customers. Experience "
                "with customer support, onboarding and stakeholder management "
                "required. CRM experience preferred."
            ),
        },
    ).json()
    res = client.post(f"{API}/cv-versions/{version['id']}/tailor", json={"jd_id": jd["id"]})
    assert res.status_code == 201, res.text
    content = res.json()["content"]
    report = res.json()["report"]

    # summary targets the role and only gains supported keywords
    assert "Customer Success Manager" in content["summary"]
    assert report["summary_keywords_added"]
    for kw in report["summary_keywords_added"]:
        assert kw.lower() in content["summary"].lower()

    # bullets re-labelled in the job's wording; facts (numbers) preserved
    assert report["bullets_relabelled"], "expected at least one re-labelled bullet"
    for r in report["bullets_relabelled"]:
        assert r["bullet"] != r["original"]
        assert re.findall(r"\d+", r["original"]) == re.findall(r"\d+", r["bullet"])
        assert r["keyword"].lower() in r["bullet"].lower()

    # role reordering flag present (bool)
    assert isinstance(report["roles_reordered"], bool)


# --- erasure ---------------------------------------------------------------------


def test_erasure_covers_new_tables(client, consented_profile, cv_id):
    v = client.post(
        f"{API}/cvs/{cv_id}/versions", json={"kind": "master_ats"}
    ).json()
    jd = client.post(
        f"{API}/profiles/{consented_profile}/job-descriptions",
        json={"title": "Customer Success Manager", "text": JD},
    ).json()
    tailored = client.post(
        f"{API}/cv-versions/{v['id']}/tailor", json={"jd_id": jd["id"]}
    ).json()

    assert client.delete(f"{API}/profiles/{consented_profile}").status_code == 204
    assert client.get(f"{API}/cv-versions/{v['id']}").status_code == 404
    assert client.get(f"{API}/tailored/{tailored['id']}").status_code == 404
    assert client.get(f"{API}/cvs/{cv_id}").status_code == 404
